import os
import csv
import docx
import pickle
from docx import Document
from datetime import date
from pprint import pprint

verbose = True
### dict lookup stuff ###

def get_lookup():
    rosetta_csv = os.path.join(script_root, "rosetta_data.csv")
    files_pickle = os.path.join(script_root, "rosetta_data.pickle")
    if os.path.exists(files_pickle):
        return pickle.load( open( files_pickle, "rb" ) )
    else:
        files = {}
        with open(rosetta_csv, encoding="utf8") as data:
            reader = csv.reader(data)
            for r in reader:
                ie, rep_pid, file_pid, file_name, pres_type, is_valid, is_well_formed, alma, tiaka, size_bytes, ext, puid = r
                files[file_pid] = (ie, rep_pid, file_pid, file_name, pres_type, is_valid, is_well_formed, alma, tiaka, size_bytes, ext, puid)
        pickle.dump( files, open( files_pickle, "wb" ) )
        return files

def make_dicts(src_folder):
    assessment_source_set = os.path.join(assessment_source_root, set_id)
    file_pids = os.listdir(src_folder)
    my_dicts = []
    for key in file_pids:
        if key not in lookup:
            print (f"Missing lookup key: {key}")
        else:
            item = (lookup[key])
            if item[7]:
                mms = f"alma:{item[7]}"
            else:
                mms = f"emu:{item[8]}"

            structured_item = {"filename":item[3],
                        "ie":item[0],
                        "mms_cms":mms,
                        "file_pid":item[2],
                        "actions":"Replace file",
                        "assessment_location":assessment_source_set+"/"+item[2]+"/"+item[3].replace(".", "_"),
                        "file_location":assessment_source_set+"/"+item[2]+"/"+item[3].replace(".", "_")+"/"+"NEW".replace("/", os.sep)}

       
            my_dicts.append(structured_item)
    return my_dicts

#########################

def setup_report_data():

	if verbose:			
		print ("Setting up report data")
	
	doc_root = os.path.join(script_root, 'reports')
	if not os.path.exists(doc_root):
		print ("Folder being used to for report location does not exist")
		print ("Please resolve before you continue")
		print (f"Folder: {doc_root}")
		quit()
	doc_name = os.path.join(doc_root, set_id+".doc")
	
	if purge_existing_report:
		if os.path.exists(doc_name):
			os.remove(doc_name)

	if os.path.exists(doc_name):
		print ()
		print ("That report has already been created at that location")
		print ("Options are delete, move, or rename the original, or change set_id if you want to carry on.")
		print ()
		print (doc_name)
		print ()
		quit()

	my_items = make_dicts(src_folder_of_files)



	boiler_plate = {"requester":"Jay Gattuso",
					"set_id":set_id, 
					"request_date":f'{date.today().strftime("%d/%m/%Y")}',
					"puid":puid,
					"format_name":format_name,
					"format_ext":format_ext,
					"error_text":error_text,
					"tool":"Rel. 1.17.49, 2017-11-02",
					"module":jhove_module}

	return doc_name, my_items, boiler_plate

#### actual report making parts ######

def add_hyperlink(paragraph, url, text, color="blue", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'single')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def get_details_from_csv(f):
	rows = []
	with open(f, encoding="utf8") as data:
		reader =  csv.reader(data,  quoting=csv.QUOTE_ALL, delimiter=',', skipinitialspace=True)
		for row in reader:
			rows.append(row)

	return rows[1:]


def is_field_flagged(field):
	flagged_fields = []
	if field in flagged_fields:
		return "True"
	else:
		return "False"

		
def make_report(set_id, my_items, boiler_plate, dest_f_name):
	if verbose:			
		print ("Starting document")

	document = Document()

	document.add_heading('File Change Request – Structural / Technical Justification')
	paragraph = document.add_paragraph('')
	requester_table = document.add_table(rows=2, cols=3)
	requester_table.style = 'Medium Shading 1 Accent 1'
	requester_table.cell(0, 0).text = "Requester"
	requester_table.cell(0, 1).text = "Set ID"
	requester_table.cell(0, 2).text = "Date"
	requester_table.cell(1, 0).text = boiler_plate["requester"]
	requester_table.cell(1, 1).text = boiler_plate["set_id"]
	requester_table.cell(1, 2).text = boiler_plate["request_date"]
	paragraph = document.add_paragraph('')

	table = document.add_table(rows=len(my_items)+1, cols=4)
	table.style = 'Medium Shading 1 Accent 1'
	table.cell(0, 0).text = "Filename"
	table.cell(0, 1).text = "IE"
	table.cell(0, 2).text = "FILE PID"
	table.cell(0, 3).text = "MMS/CMS ID"
	for i, item in enumerate(my_items, 1):
		table.cell(i, 0).text = item["filename"]
		table.cell(i, 1).text = item["ie"]
		table.cell(i, 2).text = item["file_pid"]
		table.cell(i, 3).text = item["mms_cms"]
	paragraph = document.add_paragraph('')

	document.add_heading("Error Text", level=2)
	paragraph = document.add_paragraph(f'Error message:\n\t {boiler_plate["error_text"]}')
	paragraph = document.add_paragraph(f'Tool / Version:\n\t {boiler_plate["tool"]}')
	paragraph = document.add_paragraph(f'Module / Verion:\n\t {boiler_plate["module"]}')
	paragraph = document.add_paragraph('')

	document.add_heading("Cohort Analysis", level=2)
	paragraph = document.add_paragraph(f'PRONOM Format: ')
	add_hyperlink(paragraph, f"http://www.nationalarchives.gov.uk/PRONOM/{boiler_plate['puid']}", boiler_plate["puid"])
	paragraph.add_run(f' - {boiler_plate["format_name"]} ({boiler_plate["format_ext"]})')
	paragraph = document.add_paragraph('')


	document.add_heading("Solutions", level=2)
	paragraph = document.add_paragraph('')

	document.add_heading("Recommendations", level=2)
	paragraph = document.add_paragraph('')

	document.add_heading("Summary of validation checks", level=2)

	with open (my_summary_file) as data:
		text =  data.read()
		detailed_text, summary_text = text.split("#### Summary of checks for set ####")
		summary_lines = [x for x in summary_text.split("\n") if x != ""]

	table = document.add_table(rows=len(my_items)+1, cols=6)
	table.style = 'Medium Shading 1 Accent 1'
	table.cell(0, 0).text = "File ID"
	table.cell(0, 1).text = "Basic MD Check"
	table.cell(0, 2).text = "Exiftool Check"
	table.cell(0, 3).text = "Format Specific Check"
	table.cell(0, 4).text = "JHOVE Check"
	table.cell(0, 5).text = "RMSe Check"
	for i, item in enumerate(summary_lines[1:], 1):
		item = [x for x in item.split("\t") if x != ""]
		if item != []:
			table.cell(i, 0).text = item[0]
			table.cell(i, 1).text = item[1]
			table.cell(i, 2).text = item[2]
			table.cell(i, 3).text = item[3]
			table.cell(i, 4).text = item[4]
			table.cell(i, 5).text = item[5]


	paragraph = document.add_paragraph('')

	document.add_heading("Process Notes", level=2)
	paragraph = document.add_paragraph('')
	add_hyperlink(paragraph, r"https://dia.cohesion.net.nz/Sites/IAC/SMD/NLSS/SYM/_layouts/15/WopiFrame.aspx?sourcedoc=%7bDFE88BBA-6004-4057-AA30-CFDD8840639A%7d&file=NDHA%20Change%20Request%20Process.docx&action=default", "Update Representation" )

	document.add_heading("Sign off", level=2)
	paragraph = document.add_paragraph('Authoriser Name:')
	paragraph = document.add_paragraph('Authoriser Unit/Role:')
	paragraph = document.add_paragraph('Authoriser Sign off:')
	paragraph = document.add_paragraph('Authoriser Sign off Date:')
	paragraph = document.add_paragraph('')

	document.add_heading("Processing", level=2)

	pseudo_table = document.add_paragraph('FILE PID\tRosetta IE\tAssessment\tReplacement\tActions\n')
	for i, item in enumerate(my_items, 1):
		assessment_location_text = item['file_pid']
		file_location_text = "File here"
		pseudo_table.add_run(f"{item['file_pid']}")
		pseudo_table.add_run("\t")
		pseudo_table.add_run(f"{item['ie']}")
		pseudo_table.add_run("\t")
		add_hyperlink(pseudo_table, item['assessment_location'], assessment_location_text)
		pseudo_table.add_run("\t")
		add_hyperlink(pseudo_table, item['file_location'], file_location_text)
		pseudo_table.add_run("\t")
		pseudo_table.add_run(item["actions"])
		pseudo_table.add_run("\n")
	paragraph = document.add_paragraph('')
	paragraph = document.add_paragraph('NDHA Processor:')
	paragraph = document.add_paragraph('Completed Date:')
	paragraph = document.add_paragraph('')


	document.add_page_break()

	document.add_heading("Appendix - All detected changes", level=2)
	all_fields_changed = []

	if verbose:			
		print ("Getting Appendix data")

	for pid in os.listdir(src_folder_of_files):
		my_pid = os.path.join(src_folder_of_files, pid)
		for file_set in os.listdir(my_pid):
			my_csv = os.path.join(my_pid, file_set, "differences.csv")
			if os.path.exists(my_csv):
				rows = get_details_from_csv(my_csv)
				for item in rows:
					all_fields_changed.append(item[1])

	if all_fields_changed != []:
		all_fields_changed = [x for x in list(set(all_fields_changed)) if x != ""]			
		all_fields_changed.sort()
		document.add_heading("Summary of detected changes", level=3)
		table = document.add_table(rows=len(all_fields_changed)+1, cols=2, )
		table.style = 'Medium Shading 1 Accent 1'
		table.cell(0, 0).text = "Field"
		table.cell(0, 1).text = "Flagged Field"

		for i, field in enumerate(all_fields_changed, 1):
			table.cell(i, 0).text = field
			table.cell(i, 1).text = is_field_flagged(field)
	if verbose:					
		print ("Adding diffs tables")

	for pid in os.listdir(src_folder_of_files):
		my_pid = os.path.join(src_folder_of_files, pid)
		for file_set in os.listdir(my_pid):
			my_csv = os.path.join(my_pid, file_set, "differences.csv")
			
			
			if os.path.exists(my_csv):
				if verbose:
					print ("\t", my_csv)
				last_char_index = file_set.rfind("_")
				f_name = file_set[:last_char_index] + "." + file_set[last_char_index+1:]
				paragraph = document.add_paragraph('')
				document.add_heading(f"{pid} | {f_name}", level=3)
				rows = get_details_from_csv(my_csv)

				table = document.add_table(rows=len(rows)+1, cols=5, )
				table.style = 'Medium Shading 1 Accent 1'
				table.cell(0, 0).text = "Tool"
				table.cell(0, 1).text = "Field"
				table.cell(0, 2).text = "in-file"
				table.cell(0, 3).text = "Out-file"
				table.cell(0, 4).text = "Flagged Field"
				
				for i, item in enumerate(rows, 1):
					table.cell(i, 0).text = item[0]
					table.cell(i, 1).text = item[1]
					if item[2] == "":
						table.cell(i, 2).text = "--Empty--"
					else:
						table.cell(i, 2).text = item[2]
					table.cell(i, 3).text = item[3]
					table.cell(i, 4).text = is_field_flagged(item[1])

					
	
	
	print ()
	print (f"Finished making: {doc_name}")
	print ()



	dest_check_folder = os.path.join(assessment_source_root, set_id).replace("/", os.sep)
	if not os.path.exists(dest_check_folder) or len(os.listdir(dest_check_folder)) !=  len(my_items):
		print ("It doesn't look like the file(s) have been moved to their reported location")
		print ("Please check the destination folder:")
		print (f"\n\t{dest_check_folder}\n")
		if not os.path.exists(dest_check_folder):
			os.makedirs(dest_check_folder)
	document.save(doc_name)

######################################

#### report variables ###
set_id = "fmt_353_12" ### must be unique to a set of files related by a common issue 
puid = "fmt/353" 


src_folder_of_files = f"E:/completed_cleans/{set_id}" ## this is where the list of file pids is collected from
format_name = "Tagged Image File Format"
error_text = "IFD out of sequence" 
format_ext = ".tif"
jhove_module = "TIFF-hul, Rel. 1.9.2 (2019-12-10)"
purge_existing_report = True



#### This is used to make the hyperlink. Use the final destination even it the files arn't there yet.  	
assessment_source_root = r"\\wlgprdfile13\dfs_shares\ndha\dps_export_prod\gattusoj\cleanup_phase_2" ### this is where the files will be found

### summary file data
my_summary_file = os.path.join(src_folder_of_files.replace(set_id, "summary_files"), set_id+".txt") 

script_root = r"\\wlgprdfile12\home$\Wellington\GattusoJ\homedata\Desktop\clean_up_part_2" ## this is because of how python is run here :( 
#########################

lookup = get_lookup()
doc_name, my_items, boiler_plate = setup_report_data()
make_report(set_id, my_items, boiler_plate, doc_name)
